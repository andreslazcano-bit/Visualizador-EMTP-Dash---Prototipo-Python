"""
============================================================================
EXPORTAR DOCUMENTO A WORD - PRESENTACIÓN JEFATURA
============================================================================
Convierte el documento Markdown a formato Word (.docx) con formato profesional
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from pathlib import Path

def create_word_document():
    """Crea el documento Word con formato profesional"""
    
    # Crear documento
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ========================================================================
    # PORTADA
    # ========================================================================
    
    # Logo/Título principal
    title = doc.add_heading('VISUALIZADOR EMTP', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(32)
    title_run.font.color.rgb = RGBColor(52, 83, 106)  # Color institucional #34536A
    
    # Subtítulo
    subtitle = doc.add_heading('Aspectos Clave para Jefatura', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.color.rgb = RGBColor(179, 90, 90)  # Color institucional #B35A5A
    
    # Información del documento
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Documento preparado para: ').bold = True
    info.add_run('Reunión de presentación con jefatura SEEMTP\n')
    info.add_run('Fecha: ').bold = True
    info.add_run('Noviembre 2025\n')
    info.add_run('Propósito: ').bold = True
    info.add_run('Definir decisiones estratégicas y operativas para puesta en producción')
    
    doc.add_page_break()
    
    # ========================================================================
    # ÍNDICE
    # ========================================================================
    
    doc.add_heading('📊 ÍNDICE DE CONTENIDOS', level=1)
    
    toc = [
        '1. Resumen Ejecutivo',
        '2. Funcionalidades Implementadas',
        '3. Decisiones Estratégicas Requeridas',
        '4. Definiciones Técnicas Necesarias',
        '5. Recursos y Coordinaciones',
        '6. Riesgos y Mitigaciones',
        '7. Plan de Implementación'
    ]
    
    for item in toc:
        p = doc.add_paragraph(item, style='List Bullet')
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # ========================================================================
    # 1. RESUMEN EJECUTIVO
    # ========================================================================
    
    doc.add_heading('1. RESUMEN EJECUTIVO', level=1)
    
    doc.add_heading('Estado Actual del Proyecto', level=2)
    
    # Tabla de estado
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Aspecto'
    hdr_cells[1].text = 'Estado'
    hdr_cells[2].text = 'Detalle'
    
    # Datos
    data = [
        ('Desarrollo', '✅ Funcional', 'Prototipo 100% operativo con datos simulados'),
        ('Stack Tecnológico', '✅ Moderno', 'Python 3.12 + Dash 2.14.2 + Plotly 5.18.0'),
        ('Arquitectura', '✅ Documentada', 'Diagramas y documentación técnica completa'),
        ('Autenticación', '✅ Implementada', 'Sistema de perfiles con bcrypt + JWT'),
        ('Visualizaciones', '✅ Completas', '7 módulos + mapas geográficos interactivos'),
        ('Conexión BD', '🟡 Pendiente TI', 'Scripts listos, requiere credenciales MINEDUC'),
        ('Producción', '⏳ Requiere decisiones', 'Ver secciones 3 y 4 de este documento')
    ]
    
    for i, (aspecto, estado, detalle) in enumerate(data, 1):
        cells = table.rows[i].cells
        cells[0].text = aspecto
        cells[1].text = estado
        cells[2].text = detalle
    
    doc.add_paragraph()
    
    # Valor del Proyecto
    doc.add_heading('Valor del Proyecto', level=2)
    
    valores = [
        ('Centralización', 'Un solo sistema para todos los datos EMTP (vs. múltiples Excel/R scripts dispersos)'),
        ('Accesibilidad', 'Dashboards interactivos accesibles desde cualquier navegador (vs. R/Shiny que requiere instalación)'),
        ('Actualización', 'Datos actualizados automáticamente cada semana (vs. actualización manual mensual)'),
        ('Escalabilidad', 'Arquitectura preparada para crecer con nuevas funcionalidades'),
        ('Seguridad', 'Control de acceso por perfiles + auditoría de uso')
    ]
    
    for valor, descripcion in valores:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{valor}: ').bold = True
        p.add_run(descripcion)
    
    doc.add_page_break()
    
    # ========================================================================
    # 2. FUNCIONALIDADES IMPLEMENTADAS
    # ========================================================================
    
    doc.add_heading('2. FUNCIONALIDADES IMPLEMENTADAS', level=1)
    
    modulos = [
        ('Matrícula EMTP', 
         'Evolución histórica, distribución por región/comuna/especialidad, análisis demográfico, tasas de retención',
         'Planificación de recursos educativos, identificación de especialidades con alta/baja demanda',
         'Coordinadores regionales, Jefatura SEEMTP'),
        
        ('Egresados EMTP',
         'Seguimiento de trayectorias post-egreso, transición a educación superior, inserción laboral',
         'Evaluar efectividad de articulación con educación superior',
         'Analistas SEEMTP, Investigadores'),
        
        ('Titulación EMTP',
         'Tasas de titulación por especialidad y región, tiempos promedio, identificación de cuellos de botella',
         'Detectar especialidades con problemas de titulación, implementar apoyos específicos',
         'Jefatura SEEMTP, Supervisores regionales'),
        
        ('Establecimientos EMTP',
         'Catastro de 1,124 establecimientos, distribución geográfica (16 regiones, 345 comunas)',
         'Planificación territorial de programas, asignación de recursos',
         'Coordinadores territoriales, Planificadores'),
        
        ('Docentes EMTP',
         'Perfil profesional (~5,000 docentes), especialidades por defecto/superávit, capacitación',
         'Planificación de programas de capacitación, detección de necesidades',
         'Recursos Humanos, Centros de Perfeccionamiento'),
        
        ('Mapas Geográficos',
         'Visualización coroplética de distribución territorial, mapa de establecimientos, filtros dinámicos',
         'Visualización rápida de desigualdades territoriales, presentaciones impactantes',
         'Todos los perfiles, especialmente jefatura'),
        
        ('Monitoreo y Seguimiento 🔒',
         'Gestión administrativa (convenios, rendiciones), Fortalecimiento EMTP (equipamiento, Red Futuro Técnico)',
         'Control financiero de proyectos, seguimiento de ejecución presupuestaria',
         'Solo Administradores (datos sensibles)')
    ]
    
    for i, (modulo, que_hace, para_que, usuarios) in enumerate(modulos, 1):
        doc.add_heading(f'{i}. {modulo}', level=2)
        
        p = doc.add_paragraph()
        p.add_run('¿Qué hace? ').bold = True
        p.add_run(que_hace)
        
        p = doc.add_paragraph()
        p.add_run('¿Para qué sirve? ').bold = True
        p.add_run(para_que)
        
        p = doc.add_paragraph()
        p.add_run('Usuarios principales: ').bold = True
        p.add_run(usuarios)
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========================================================================
    # 3. DECISIONES ESTRATÉGICAS
    # ========================================================================
    
    doc.add_heading('3. DECISIONES ESTRATÉGICAS REQUERIDAS', level=1)
    
    # DECISIÓN 1: Modelo de Acceso
    doc.add_heading('DECISIÓN 1: Modelo de Acceso y Usuarios', level=2)
    
    p = doc.add_paragraph()
    p.add_run('⭐ RECOMENDACIÓN: Opción B - Perfiles con Login').bold = True
    p_format = p.paragraph_format
    p_format.space_after = Pt(12)
    
    doc.add_paragraph('Razones:')
    razones = [
        'Sección "Monitoreo y Seguimiento de Proyectos" contiene datos sensibles (convenios, rendiciones)',
        'Auditoría es crítica para saber quién accede a qué información',
        'Escalabilidad: Permite agregar más perfiles en el futuro',
        'Compliance: Cumplimiento de normativas de protección de datos'
    ]
    for razon in razones:
        doc.add_paragraph(razon, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Perfiles propuestos:')
    
    # Tabla de perfiles
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Perfil'
    hdr_cells[1].text = 'Usuarios Típicos'
    hdr_cells[2].text = 'Permisos'
    hdr_cells[3].text = 'Cantidad'
    
    perfiles_data = [
        ('👤 Usuario Básico', 'Directores, Docentes', 'Ver dashboards públicos, mapas\n❌ Sin acceso a Proyectos', '~100-200'),
        ('👔 Analista SEEMTP', 'Coordinadores regionales', 'Todo lo anterior + Exportar datos + Filtros avanzados', '~30-50'),
        ('⚙️ Administrador', 'Jefatura SEEMTP, TI', 'Acceso total + Gestión de usuarios + Proyectos', '~5-10')
    ]
    
    for i, (perfil, usuarios, permisos, cantidad) in enumerate(perfiles_data, 1):
        cells = table.rows[i].cells
        cells[0].text = perfil
        cells[1].text = usuarios
        cells[2].text = permisos
        cells[3].text = cantidad
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('❓ PREGUNTA CLAVE: ').bold = True
    p.add_run('¿Están de acuerdo con este modelo de perfiles?')
    
    doc.add_paragraph()
    
    # DECISIÓN 2: Autenticación
    doc.add_heading('DECISIÓN 2: Método de Autenticación', level=2)
    
    p = doc.add_paragraph()
    p.add_run('⭐ RECOMENDACIÓN: Microsoft 365 / Entra ID').bold = True
    
    doc.add_paragraph('Razones:')
    razones_auth = [
        'Todos los funcionarios MINEDUC ya tienen cuenta Microsoft 365',
        'Experiencia de usuario fluida: "Iniciar sesión con Microsoft" (un click)',
        'Seguridad robusta: MFA, políticas de contraseñas institucionales',
        'Sin gestión manual: TI no tiene que crear/desactivar cuentas',
        'Estándar moderno: OAuth 2.0 es el estándar de la industria'
    ]
    for razon in razones_auth:
        doc.add_paragraph(razon, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('❓ PREGUNTA CLAVE: ').bold = True
    p.add_run('¿Prefieren integración con Microsoft 365 o gestionar credenciales propias?')
    
    doc.add_paragraph()
    
    # DECISIÓN 3: Alcance
    doc.add_heading('DECISIÓN 3: Alcance de Funcionalidades', level=2)
    
    p = doc.add_paragraph()
    p.add_run('⭐ RECOMENDACIÓN: Visualización + Reportería Básica (Excel/PDF)').bold = True
    
    doc.add_paragraph('Razones:')
    razones_alcance = [
        'Analistas necesitan compartir datos en reuniones, informes, presentaciones',
        'Excel es el formato estándar de trabajo en el ministerio',
        'PDF para reportes ejecutivos con gráficos incluidos',
        'Equilibrio esfuerzo/beneficio: Gran valor con desarrollo moderado (+2 semanas)'
    ]
    for razon in razones_alcance:
        doc.add_paragraph(razon, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('❓ PREGUNTA CLAVE: ').bold = True
    p.add_run('¿Es suficiente con exportación básica o necesitan reportes automatizados?')
    
    doc.add_paragraph()
    
    # DECISIÓN 4: Actualización
    doc.add_heading('DECISIÓN 4: Frecuencia de Actualización de Datos', level=2)
    
    p = doc.add_paragraph()
    p.add_run('⭐ RECOMENDACIÓN: Actualización Semanal (cada lunes 2AM)').bold = True
    
    doc.add_paragraph('Razones:')
    razones_actualizacion = [
        'Datos educativos NO cambian minuto a minuto (matrícula, titulación son anuales/semestrales)',
        'SIGE se actualiza semanalmente → Sincronizamos después de su actualización',
        'Dashboards instantáneos: 0.5 segundos vs. 5-10 segundos con SQL directo',
        'Sin sobrecarga de bases de datos productivas de MINEDUC',
        'Sistema ya implementado 100% (scripts listos)'
    ]
    for razon in razones_actualizacion:
        doc.add_paragraph(razon, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('❓ PREGUNTA CLAVE: ').bold = True
    p.add_run('¿Es aceptable que los datos tengan máximo 7 días de antigüedad?')
    
    doc.add_page_break()
    
    # ========================================================================
    # 4. DEFINICIONES TÉCNICAS
    # ========================================================================
    
    doc.add_heading('4. DEFINICIONES TÉCNICAS NECESARIAS', level=1)
    
    # Hosting
    doc.add_heading('Infraestructura de Hosting', level=2)
    
    p = doc.add_paragraph()
    p.add_run('⭐ RECOMENDACIÓN: Azure App Service').bold = True
    
    doc.add_paragraph('Razones:')
    razones_azure = [
        'Ecosistema Microsoft: MINEDUC ya usa M365, Teams, SharePoint → sinergia',
        'Integración nativa con Entra ID para autenticación',
        'Compliance chileno: Azure tiene datacenter en Brasil (latencia baja)',
        'Soporte técnico Microsoft incluido',
        'Escalamiento automático: Si usuarios crecen, el servidor se adapta'
    ]
    for razon in razones_azure:
        doc.add_paragraph(razon, style='List Bullet')
    
    doc.add_paragraph()
    
    # Tabla de costos
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Opción'
    hdr_cells[1].text = 'Costo Anual'
    
    cells = table.rows[1].cells
    cells[0].text = 'Azure App Service (Recomendado)'
    cells[1].text = '~$1,200 USD/año'
    
    cells = table.rows[2].cells
    cells[0].text = 'Servidor On-Premise MINEDUC'
    cells[1].text = '$0 (usa infraestructura existente)'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Comparación vs. alternativas comerciales:').bold = True
    
    comparaciones = [
        'Licencia Tableau: ~$42,000 USD/año (50 usuarios)',
        'Licencia Power BI Pro: ~$6,000 USD/año (50 usuarios)',
        'Esta solución: ~$1,200 USD/año = 95% más barato que Power BI ✅'
    ]
    for comp in comparaciones:
        doc.add_paragraph(comp, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('❓ PREGUNTA CLAVE: ').bold = True
    p.add_run('¿Tienen presupuesto para hosting cloud (~$600 USD/año) o prefieren servidor interno?')
    
    doc.add_paragraph()
    
    # Conexión a Bases de Datos
    doc.add_heading('🔴 MÁS URGENTE: Conexión a Bases de Datos MINEDUC', level=2)
    
    p = doc.add_paragraph()
    p.add_run('ACCIÓN REQUERIDA: ').bold = True
    p.add_run('Programar reunión con Jefe TI MINEDUC ASAP')
    
    doc.add_paragraph()
    doc.add_paragraph('Bases de datos requeridas:')
    
    bases = [
        'SIGE: Matrícula EMTP, Establecimientos, Cursos (Actualización: Semanal)',
        'Sistema de Titulados: Titulación por especialidad (Actualización: Mensual)',
        'Sistema Financiero: Convenios, Rendiciones, Presupuesto (Actualización: Semanal)'
    ]
    for base in bases:
        doc.add_paragraph(base, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Lo que necesitamos de TI:')
    
    necesidades = [
        'Credenciales de acceso SQL Server (solo lectura)',
        'Whitelist de IP del servidor de la app',
        'Reglas de firewall para puerto 1433',
        'VPN si es necesario (acceso desde Azure)'
    ]
    for necesidad in necesidades:
        doc.add_paragraph(necesidad, style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('✅ Scripts ya implementados ').bold = True
    p.add_run('(listos para usar cuando TI entregue credenciales):')
    
    scripts = [
        'scripts/test_connections.py - Verificar conectividad',
        'scripts/actualizar_datos_semanal.py - Actualización automática',
        'src/data/loaders.py - Cargador de datos con cache'
    ]
    for script in scripts:
        doc.add_paragraph(script, style='List Bullet')
    
    doc.add_page_break()
    
    # ========================================================================
    # 5. RECURSOS Y COORDINACIONES
    # ========================================================================
    
    doc.add_heading('5. RECURSOS Y COORDINACIONES', level=1)
    
    doc.add_heading('Equipo Necesario', level=2)
    
    # Tabla de equipo
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rol'
    hdr_cells[1].text = 'Responsabilidad'
    hdr_cells[2].text = 'Dedicación'
    hdr_cells[3].text = '¿Quién?'
    
    equipo_data = [
        ('Líder de Proyecto', 'Decisiones estratégicas, priorización', '20% (1 día/semana)', 'Jefatura SEEMTP'),
        ('Desarrollador Principal', 'Desarrollo, mantenimiento, bugs', '100% (1-2 meses)', 'Actual (Andrés)'),
        ('TI MINEDUC', 'Accesos BD, infraestructura', '20% (durante setup)', 'Coordinador TI'),
        ('Analista Funcional', 'Validación de datos, pruebas', '10% (durante desarrollo)', 'Analista SEEMTP')
    ]
    
    for i, (rol, resp, dedic, quien) in enumerate(equipo_data, 1):
        cells = table.rows[i].cells
        cells[0].text = rol
        cells[1].text = resp
        cells[2].text = dedic
        cells[3].text = quien
    
    doc.add_paragraph()
    
    doc.add_heading('Presupuesto Estimado', level=2)
    
    # Tabla de presupuesto
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ítem'
    hdr_cells[1].text = 'Costo'
    hdr_cells[2].text = 'Frecuencia'
    
    presupuesto_data = [
        ('Desarrollo inicial', '$0', 'Una vez (ya realizado)'),
        ('Hosting Azure', '$50-100 USD', 'Mensual'),
        ('Base de datos Azure SQL', '$30-50 USD', 'Mensual (si no usan on-premise)'),
        ('Dominio personalizado', '$12 USD', 'Anual (opcional)'),
        ('Certificado SSL', '$0', 'Incluido en Azure'),
        ('TOTAL ANUAL (cloud)', '~$1,200 USD/año', 'Escenario Azure')
    ]
    
    for i, (item, costo, freq) in enumerate(presupuesto_data, 1):
        cells = table.rows[i].cells
        cells[0].text = item
        cells[1].text = costo
        cells[2].text = freq
    
    doc.add_page_break()
    
    # ========================================================================
    # 6. RIESGOS Y MITIGACIONES
    # ========================================================================
    
    doc.add_heading('6. RIESGOS Y MITIGACIONES', level=1)
    
    riesgos = [
        ('Retraso en Acceso a Bases de Datos',
         '🟡 Media',
         '🔴 Alto (bloquea conexión a datos reales)',
         [
             'Iniciar coordinación con TI de inmediato (memo ya preparado)',
             'Mientras tanto, seguir trabajando con datos simulados',
             'Pedir acceso solo de lectura (menos burocracia)'
         ]),
        
        ('Sobrecarga de Usuarios (Escalabilidad)',
         '🟢 Baja',
         '🟡 Medio (sistema lento)',
         [
             'Azure autoscaling: Servidor crece automáticamente',
             'Cache de datos: Dashboards cargan desde archivos locales',
             'Monitoreo: Alertas si tiempo de respuesta >3 segundos'
         ]),
        
        ('Resistencia al Cambio (Usuarios prefieren Excel/R)',
         '🟡 Media',
         '🟡 Medio (baja adopción)',
         [
             'Capacitación inicial: Sesión de 1 hora mostrando beneficios',
             'Usuarios piloto: Seleccionar "early adopters" entusiastas',
             'Exportación a Excel: Permitir llevar datos a herramienta favorita',
             'No forzar: Sistema coexiste con Excel/R'
         ])
    ]
    
    for riesgo, prob, impacto, mitigaciones in riesgos:
        doc.add_heading(f'⚠️ {riesgo}', level=2)
        
        p = doc.add_paragraph()
        p.add_run('Probabilidad: ').bold = True
        p.add_run(prob)
        
        p = doc.add_paragraph()
        p.add_run('Impacto: ').bold = True
        p.add_run(impacto)
        
        p = doc.add_paragraph()
        p.add_run('Mitigación:').bold = True
        
        for mitigacion in mitigaciones:
            doc.add_paragraph(mitigacion, style='List Bullet')
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========================================================================
    # 7. PLAN DE IMPLEMENTACIÓN
    # ========================================================================
    
    doc.add_heading('7. PLAN DE IMPLEMENTACIÓN', level=1)
    
    doc.add_heading('Cronograma Estimado: 12 Semanas', level=2)
    
    fases = [
        ('Fase 1: Definiciones y Coordinación (Semanas 1-2)',
         [
             'Reunión de presentación a jefatura',
             'Elaborar memo para Jefe TI MINEDUC',
             'Definir usuarios piloto (3-5 personas)',
             'Aprobar presupuesto (si es Azure)'
         ],
         'Decisiones estratégicas tomadas + Coordinación TI iniciada'),
        
        ('Fase 2: Configuración Técnica (Semanas 3-5)',
         [
             'TI entrega credenciales SQL Server',
             'Configurar servidor (Azure o on-premise)',
             'Probar conexiones a bases de datos',
             'Primera actualización de datos reales',
             'Configurar autenticación M365 (si aprobado)'
         ],
         'Sistema conectado a datos reales de MINEDUC'),
        
        ('Fase 3: Desarrollo de Funcionalidades (Semanas 6-8)',
         [
             'Implementar exportación a Excel',
             'Implementar exportación a PDF',
             'Agregar filtros avanzados para analistas',
             'Validación con usuarios piloto'
         ],
         'Funcionalidades de reportería implementadas'),
        
        ('Fase 4: Testing y Refinamiento (Semanas 9-10)',
         [
             'Testing de carga (simular 50 usuarios)',
             'Revisión de usabilidad (usuarios piloto)',
             'Validación de datos vs. fuentes oficiales',
             'Documentación de usuario final',
             'Configurar monitoreo y alertas'
         ],
         'Sistema validado y refinado'),
        
        ('Fase 5: Capacitación y Lanzamiento (Semanas 11-12)',
         [
             'Sesión de capacitación a usuarios finales (1h)',
             'Crear usuarios en sistema',
             'Comunicación oficial de lanzamiento',
             'Soporte "hot" primera semana',
             'Monitoreo de adopción'
         ],
         '🎉 Sistema en producción y usuarios activos')
    ]
    
    for fase, tareas, hito in fases:
        doc.add_heading(fase, level=3)
        
        doc.add_paragraph('Tareas:')
        for tarea in tareas:
            doc.add_paragraph(tarea, style='List Bullet')
        
        p = doc.add_paragraph()
        p.add_run('Hito: ').bold = True
        p.add_run(hito)
        
        doc.add_paragraph()
    
    doc.add_paragraph()
    
    # Indicadores de éxito
    doc.add_heading('Indicadores de Éxito (Post-Lanzamiento)', level=2)
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Indicador'
    hdr_cells[1].text = 'Meta'
    
    indicadores = [
        ('Adopción', '>60% de usuarios invitados acceden en primer mes'),
        ('Uso recurrente', '>30% de usuarios acceden semanalmente'),
        ('Satisfacción', '>80% de usuarios lo encuentran útil'),
        ('Performance', 'Dashboards cargan en <3 segundos'),
        ('Exportaciones', '>50 reportes exportados en primer mes'),
        ('Disponibilidad', '>99% uptime (máximo 7 horas caídas/mes)')
    ]
    
    for i, (indicador, meta) in enumerate(indicadores, 1):
        cells = table.rows[i].cells
        cells[0].text = indicador
        cells[1].text = meta
    
    doc.add_page_break()
    
    # ========================================================================
    # 8. CHECKLIST DE DECISIONES
    # ========================================================================
    
    doc.add_heading('8. CHECKLIST DE DECISIONES', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Usar en la reunión para marcar decisiones tomadas:').italic = True
    
    doc.add_paragraph()
    
    doc.add_heading('Decisiones Estratégicas', level=2)
    
    decisiones_estrategicas = [
        'Modelo de acceso: Público / Con perfiles / Híbrido → Decisión: __________',
        'Autenticación: AD / Microsoft 365 / Credenciales propias → Decisión: __________',
        'Alcance funcionalidades: Solo viz / +Reportería / Plataforma completa → Decisión: __________',
        'Actualización datos: Tiempo real / Diaria / Semanal / Mensual → Decisión: __________'
    ]
    
    for decision in decisiones_estrategicas:
        doc.add_paragraph(decision, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Definiciones Técnicas', level=2)
    
    decisiones_tecnicas = [
        'Hosting: On-premise / Azure / AWS / Heroku → Decisión: __________',
        'Presupuesto aprobado: Sí / No / Pendiente → Decisión: __________',
        'Notificaciones: Sí / No / Solo errores → Decisión: __________'
    ]
    
    for decision in decisiones_tecnicas:
        doc.add_paragraph(decision, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Coordinaciones', level=2)
    
    coordinaciones = [
        'Reunión con TI programada: Fecha: __________',
        'Usuarios piloto seleccionados: Nombres: __________',
        'Fecha tentativa de lanzamiento: __________'
    ]
    
    for coord in coordinaciones:
        doc.add_paragraph(coord, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Última actualización: 17 de noviembre de 2025').italic = True
    
    # ========================================================================
    # GUARDAR DOCUMENTO
    # ========================================================================
    
    output_path = Path(__file__).parent.parent / 'docs' / 'PRESENTACION_JEFATURA_ASPECTOS_CLAVE.docx'
    doc.save(output_path)
    
    print(f"✅ Documento Word generado exitosamente:")
    print(f"📄 {output_path}")
    print(f"📊 Tamaño: {output_path.stat().st_size / 1024:.2f} KB")
    
    return output_path


if __name__ == '__main__':
    create_word_document()
